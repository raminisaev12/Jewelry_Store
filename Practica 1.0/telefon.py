import flet as ft
import pymysql
import pymysql.cursors
from datetime import datetime, date
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
import os
import tempfile

# ---------- Подключение к БД (Clever Cloud) ----------
def connect_to_db():
    try:
        connection = pymysql.connect(
            host="blc7gqdzfqspj5pimbbe-mysql.services.clever-cloud.com",
            port=3306,
            user="u0bp5lienenha2l1",
            password="YHsQ0s53cdT4e2lVChNF",
            database="blc7gqdzfqspj5pimbbe",
            cursorclass=pymysql.cursors.Cursor
        )
        print("DEBUG: Успешно! Подключились к Clever Cloud")
        return connection
    except Exception as e:
        print(f"Ошибка подключения к БД: {e}")
        return None

def get_distinct_values(column):
    conn = connect_to_db()
    if not conn:
        return []
    try:
        cursor = conn.cursor()
        cursor.execute(
            f"SELECT DISTINCT {column} FROM jewelry_items WHERE {column} IS NOT NULL AND {column} != '' ORDER BY {column}")
        rows = cursor.fetchall()
        return [r[0] for r in rows]
    except Exception as e:
        print(e)
        return []
    finally:
        conn.close()

def main(page: ft.Page):
    page.title = "Алмазный путь"
    page.theme_mode = ft.ThemeMode.LIGHT
    page.bgcolor = "#FDF5E6"
    page.padding = 10
    page.scroll = ft.ScrollMode.AUTO

    page.window_width = 390
    page.window_height = 844
    page.window_resizable = True

    page.data = {"current_sort": "", "selected_item": None}

    btn_color = "#2D6A4F"
    danger_color = "#8B0000"


    status_msg = ft.Text("", size=14, weight="bold", color="red")
    selected_label = ft.Text("", size=14, weight="bold", color="blue")
    buy_status_msg = ft.Text("", size=14, weight="bold", color="red")

    def show_message(text, color="red", target="form"):
        if target == "form":
            status_msg.value = text
            status_msg.color = color
        else:
            buy_status_msg.value = text
            buy_status_msg.color = color
        page.update()

    # ---------- Поля ввода ----------
    f_name = ft.TextField(label="Название изделия", expand=True,
                          input_filter=ft.InputFilter(allow=True, regex_string=r"[А-Яа-яA-Za-z\s\-]+"))
    f_type = ft.Dropdown(
        label="Тип",
        options=[ft.dropdown.Option(x) for x in ["Кольцо", "Серьги", "Браслет", "Цепочка"]],
        value="Кольцо", expand=True
    )
    f_cat = ft.Dropdown(
        label="Категория",
        options=[ft.dropdown.Option(x) for x in ["Свадебные", "Женские", "Мужские", "Детские"]],
        expand=True
    )
    f_metal = ft.Dropdown(
        label="Металл",
        options=[ft.dropdown.Option(x) for x in
                 ["Золото", "Красное золото", "Белое золото", "Желтое золото", "Серебро"]],
        expand=True
    )
    f_gem = ft.Dropdown(
        label="Камень",
        options=[ft.dropdown.Option(x) for x in ["Нет", "Бриллиант", "Сапфир", "Рубин", "Изумруд", "Фианит"]],
        expand=True
    )
    f_purity = ft.Dropdown(
        label="Проба",
        options=[ft.dropdown.Option(x) for x in ["375", "500", "585", "750", "925", "958", "999"]],
        expand=True
    )
    f_weight = ft.TextField(label="Вес (г)", keyboard_type=ft.KeyboardType.NUMBER, expand=True,
                            input_filter=ft.InputFilter(allow=True, regex_string=r"\d*\.?\d*"))
    f_price = ft.TextField(label="Цена (₽)", keyboard_type=ft.KeyboardType.NUMBER, expand=True,
                           input_filter=ft.InputFilter(allow=True, regex_string=r"\d*\.?\d*"))
    f_size = ft.Dropdown(label="Размер", options=[], expand=True)
    f_quantity = ft.TextField(label="Количество", keyboard_type=ft.KeyboardType.NUMBER, value="1", expand=True,
                              input_filter=ft.InputFilter(allow=True, regex_string=r"\d*"))

    def update_size_options(e=None):
        selected = f_type.value
        if selected == "Кольцо":
            sizes = ["15", "15.5", "16", "16.5", "17", "17.5", "18", "18.5", "19"]
        elif selected == "Браслет":
            sizes = ["16", "17", "18", "19", "20"]
        elif selected == "Цепочка":
            sizes = ["40", "45", "50", "55", "60", "65"]
        else:
            sizes = ["—"]
        f_size.options = [ft.dropdown.Option(s) for s in sizes]
        f_size.value = sizes[0]
        page.update()

    update_size_options()

    def clear_form():
        f_name.value = ""
        f_weight.value = ""
        f_price.value = ""
        f_quantity.value = "1"
        f_purity.value = None
        f_cat.value = None
        f_metal.value = None
        f_gem.value = None
        f_type.value = "Кольцо"
        update_size_options()

    # ---------- Таблица товаров ----------
    data_table = ft.DataTable(
        columns=[
            ft.DataColumn(ft.Text("ID")),
            ft.DataColumn(ft.Text("Название")),
            ft.DataColumn(ft.Text("Тип")),
            ft.DataColumn(ft.Text("Категория")),
            ft.DataColumn(ft.Text("Металл")),
            ft.DataColumn(ft.Text("Камень")),
            ft.DataColumn(ft.Text("Проба")),
            ft.DataColumn(ft.Text("Вес")),
            ft.DataColumn(ft.Text("Размер")),
            ft.DataColumn(ft.Text("Кол-во")),
            ft.DataColumn(ft.Text("Цена")),
        ],
        rows=[],
        column_spacing=15,
        heading_row_height=40, data_row_min_height=35,
        heading_text_style=ft.TextStyle(size=12, weight="bold"),
        data_text_style=ft.TextStyle(size=11)
    )

    # ---------- Поиск и фильтры ----------
    search_field = ft.TextField(label="Поиск по названию изделия...", expand=True)
    filter_type = ft.Dropdown(label="Тип", width=110)
    filter_cat = ft.Dropdown(label="Категория", width=110)
    filter_metal = ft.Dropdown(label="Металл", width=110)
    filter_gem = ft.Dropdown(label="Камень", width=110)
    filter_size = ft.Dropdown(label="Размер", width=110)
    filter_purity = ft.Dropdown(label="Проба", width=110)

    def load_filter_values():
        filter_type.options = [ft.dropdown.Option("Все")] + [ft.dropdown.Option(x) for x in get_distinct_values("type")]
        filter_cat.options = [ft.dropdown.Option("Все")] + [ft.dropdown.Option(x) for x in
                                                            get_distinct_values("category")]
        filter_metal.options = [ft.dropdown.Option("Все")] + [ft.dropdown.Option(x) for x in
                                                              get_distinct_values("metal")]
        filter_gem.options = [ft.dropdown.Option("Все")] + [ft.dropdown.Option(x) for x in
                                                            get_distinct_values("gemstone")]
        filter_size.options = [ft.dropdown.Option("Все")] + [ft.dropdown.Option(x) for x in get_distinct_values("size")]
        filter_purity.options = [ft.dropdown.Option("Все")] + [ft.dropdown.Option(x) for x in
                                                               get_distinct_values("purity")]
        for dd in [filter_type, filter_cat, filter_metal, filter_gem, filter_size, filter_purity]:
            dd.value = "Все"
        page.update()

    load_filter_values()

    # ---------- Функция выбора товара ----------
    def select_item(e, row_data):
        page.data["selected_item"] = row_data
        selected_label.value = f"Выбран: {row_data[0]} {row_data[1]}"
        selected_label.color = "blue"
        confirm_delete_panel.visible = False
        buy_status_msg.value = ""  # очищаем статус
        page.update()

    def load_data():
        data_table.rows.clear()
        search = search_field.value
        filters = {
            "type": filter_type.value if filter_type.value != "Все" else None,
            "category": filter_cat.value if filter_cat.value != "Все" else None,
            "metal": filter_metal.value if filter_metal.value != "Все" else None,
            "gemstone": filter_gem.value if filter_gem.value != "Все" else None,
            "size": filter_size.value if filter_size.value != "Все" else None,
            "purity": filter_purity.value if filter_purity.value != "Все" else None,
        }
        conn = connect_to_db()
        if not conn:
            show_message("Ошибка подключения к БД", target="form")
            return
        try:
            cursor = conn.cursor()
            query = "SELECT CONCAT(prefix, '-', id), name, type, category, metal, gemstone, purity, weight, size, quantity, price FROM jewelry_items WHERE 1=1"
            params = []
            if search:
                query += " AND name LIKE %s"
                params.append(f"%{search}%")
            for col, val in filters.items():
                if val:
                    query += f" AND {col} = %s"
                    params.append(val)
            if page.data["current_sort"]:
                query += f" ORDER BY {page.data['current_sort']}"
            cursor.execute(query, params)
            for row in cursor.fetchall():
                id_button = ft.TextButton(
                    content=ft.Text(str(row[0])),
                    on_click=lambda e, r=row: select_item(e, r),
                    style=ft.ButtonStyle(color=ft.Colors.BLUE)
                )
                cells = [ft.DataCell(id_button)]
                for val in row[1:]:
                    cells.append(ft.DataCell(ft.Text(str(val))))
                data_table.rows.append(ft.DataRow(cells=cells, data=row))
            cursor.close()
            conn.close()
        except Exception as e:
            print(f"Ошибка загрузки: {e}")
        page.update()

    search_field.on_change = lambda _: load_data()
    for dd in [filter_type, filter_cat, filter_metal, filter_gem, filter_size, filter_purity]:
        dd.on_change = lambda _: load_data()

    # ---------- Добавление ----------
    def add_item(e):
        show_message("Добавление товара...", "blue", target="form")
        if not f_name.value or not f_price.value:
            show_message("Ошибка: Заполните название и цену изделия!", target="form")
            return
        try:
            weight_val = float(f_weight.value) if f_weight.value else 0.0
            price_val = float(f_price.value) if f_price.value else 0.0
            qty_val = int(f_quantity.value) if f_quantity.value.isdigit() else 1
        except ValueError:
            show_message("Ошибка: Вес, Цена и Количество должны быть числами!", target="form")
            return
        conn = connect_to_db()
        if not conn:
            show_message("Ошибка: Нет связи с базой данных!", target="form")
            return
        try:
            cursor = conn.cursor()
            prefix_map = {
                "Кольцо":  ("A", 1000, 1999),
                "Серьги":  ("B", 2000, 2999),
                "Браслет": ("C", 3000, 3999),
                "Цепочка": ("D", 4000, 4999)
            }
            prefix, min_id, max_id = prefix_map.get(f_type.value, ("A", 1000, 1999))
            cursor.execute(
                "SELECT MAX(id) FROM jewelry_items WHERE prefix = %s AND id BETWEEN %s AND %s",
                (prefix, min_id, max_id)
            )
            row = cursor.fetchone()
            if row[0] is None:
                new_id = min_id
            else:
                if row[0] >= max_id:
                    show_message(f"Достигнут лимит ID для типа «{f_type.value}» (макс. {max_id})", target="form")
                    conn.close()
                    return
                new_id = row[0] + 1
            cursor.execute("""
                INSERT INTO jewelry_items (id, prefix, name, type, category, metal, gemstone, purity, weight, size, quantity, price)
                VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
            """, (new_id, prefix, f_name.value, f_type.value, f_cat.value, f_metal.value,
                  f_gem.value, f_purity.value, weight_val, f_size.value, qty_val, price_val))
            conn.commit()
            cursor.close()
            conn.close()
            clear_form()
            show_message(f"Изделие успешно добавлено! (ID: {prefix}-{new_id})", "green", target="form")
            load_data()
        except Exception as ex:
            show_message(f"Ошибка базы данных: {ex}", target="form")
            print(f"Ошибка добавления: {ex}")

    # ---------- Встроенная форма покупки ----------
    buy_form = ft.Container(visible=False, padding=10)

    seller_dd = ft.Dropdown(
        label="Продавец ФИО:",
        options=[ft.dropdown.Option(x) for x in ["Иванов И.И.", "Петров П.П.", "Сидорова А.А."]],
        value="Иванов И.И.", width=200
    )
    qty_field = ft.TextField(label="Количество:", value="1",
                             keyboard_type=ft.KeyboardType.NUMBER,
                             input_filter=ft.InputFilter(allow=True, regex_string=r"\d*"))
    payment_dd = ft.Dropdown(
        label="Способ оплаты",
        options=[ft.dropdown.Option(x) for x in ["Наличные", "Карта", "Перевод"]],
        value="Наличные", width=200
    )
    total_text = ft.Text("", size=16, color="green")

    def open_buy_form(e):
        selected = page.data["selected_item"]
        if not selected:
            show_message("Ошибка: Сначала выберите товар!", target="buy")
            return
        price = float(selected[10])
        stock = int(selected[9])
        qty_field.value = "1"
        total_text.value = f"К оплате: {price:.2f} ₽"
        buy_form.data = selected
        buy_form.visible = True
        confirm_delete_panel.visible = False
        buy_status_msg.value = ""   # очищаем
        page.update()

    def recalc_total(*args):
        selected = buy_form.data
        if selected:
            price = float(selected[10])
            qty = int(qty_field.value) if qty_field.value.isdigit() else 0
            total = qty * price
            total_text.value = f"К оплате: {total:.2f} ₽"
            page.update()

    qty_field.on_change = recalc_total

    def confirm_purchase(e):
        selected = buy_form.data
        if not selected:
            return
        full_id = str(selected[0])
        item_name = selected[1]
        price = float(selected[10])
        stock = int(selected[9])
        numeric_id = full_id.split("-")[1] if "-" in full_id else full_id
        qty = int(qty_field.value) if qty_field.value.isdigit() else 0
        if qty <= 0 or qty > stock:
            show_message("Неверное количество!", "red", target="buy")
            return
        total = qty * price
        try:
            conn = connect_to_db()
            cursor = conn.cursor()
            now = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            cursor.execute("UPDATE jewelry_items SET quantity = quantity - %s WHERE id = %s", (qty, numeric_id))
            cursor.execute("""
                INSERT INTO sales_history (seller_name, item_name, quantity_sold, price_per_item, payment_method, total_price, sale_date)
                VALUES (%s,%s,%s,%s,%s,%s,%s)
            """, (seller_dd.value, item_name, qty, price, payment_dd.value, total, now))
            conn.commit()
            cursor.close()
            conn.close()
            buy_form.visible = False
            show_message(f"Покупка оформлена ({item_name})!", "green", target="buy")
            load_data()
        except Exception as ex:
            show_message(f"Ошибка продажи: {ex}", target="buy")

    buy_form.content = ft.Column([
        ft.Text("Оформление продажи", weight="bold"),
        seller_dd,
        qty_field,
        payment_dd,
        total_text,
        ft.Row([
            ft.Button("Отмена", on_click=lambda _: setattr(buy_form, 'visible', False) or page.update()),
            ft.Button("Подтвердить", on_click=confirm_purchase)
        ])
    ])

    # ---------- Удаление с панелью подтверждения ----------
    confirm_delete_panel = ft.Container(visible=False, padding=10)

    def show_delete_confirm(e):
        if not page.data["selected_item"]:
            show_message("Ошибка: Сначала выберите товар!", target="buy")
            return
        confirm_delete_panel.visible = True
        buy_form.visible = False
        buy_status_msg.value = ""
        page.update()

    def cancel_delete(e):
        confirm_delete_panel.visible = False
        page.update()

    def execute_delete(e):
        confirm_delete_panel.visible = False
        selected = page.data["selected_item"]
        if not selected:
            return
        full_id = str(selected[0])
        numeric_id = full_id.split("-")[1] if "-" in full_id else full_id
        conn = connect_to_db()
        if not conn:
            return
        try:
            cursor = conn.cursor()
            cursor.execute("DELETE FROM jewelry_items WHERE id = %s", (numeric_id,))
            conn.commit()
            cursor.close()
            conn.close()
            page.data["selected_item"] = None
            selected_label.value = ""
            show_message("Изделие успешно удалено!", "green", target="buy")
            load_data()
        except Exception as ex:
            show_message(f"Ошибка удаления: {ex}", target="buy")

    confirm_delete_panel.content = ft.Column([
        ft.Text("Вы уверены, что хотите удалить изделие?", weight="bold"),
        ft.Row([
            ft.Button("Да", on_click=execute_delete, bgcolor="red", color="white"),
            ft.Button("Нет", on_click=cancel_delete)
        ])
    ])

    delete_btn = ft.Button("Удалить", bgcolor=danger_color, color="white", on_click=show_delete_confirm)

    # ---------- Сортировка ----------
    sort_button = ft.PopupMenuButton(
        content=ft.Text("Сортировки"),
        items=[
            ft.PopupMenuItem(content=ft.Text("Сначала дешевые"),
                             on_click=lambda _: (page.data.update({"current_sort": "price ASC"}), load_data())),
            ft.PopupMenuItem(content=ft.Text("Сначала дорогие"),
                             on_click=lambda _: (page.data.update({"current_sort": "price DESC"}), load_data())),
            ft.PopupMenuItem(content=ft.Text("Сортировка от А до Я"),
                             on_click=lambda _: (page.data.update({"current_sort": "name ASC"}), load_data())),
            ft.PopupMenuItem(content=ft.Text("Сортировка от Я до А"),
                             on_click=lambda _: (page.data.update({"current_sort": "name DESC"}), load_data())),
            ft.PopupMenuItem(),
            ft.PopupMenuItem(content=ft.Text("Сбросить сортировку"),
                             on_click=lambda _: (page.data.update({"current_sort": ""}), load_data())),
        ]
    )

    # ---------- Отчёты (с кнопками-календарями) ----------
    report_date_from = ft.TextField(label="С:", width=100, hint_text="ГГГГ-ММ-ДД", read_only=True)
    report_date_to = ft.TextField(label="По:", width=100, hint_text="ГГГГ-ММ-ДД", read_only=True)

    def open_date_picker(for_field):
        def on_date_selected(e):
            if e.control.value:
                for_field.value = e.control.value.strftime('%Y-%m-%d')
            else:
                for_field.value = ""
            page.update()

        picker = ft.DatePicker(
            on_change=on_date_selected,
            first_date=datetime(2020, 1, 1),
            last_date=datetime(2030, 12, 31)
        )
        page.overlay.append(picker)
        picker.open = True
        page.update()

    report_seller = ft.Dropdown(label="Продавец", options=[ft.dropdown.Option(x) for x in
                                                           ["— Все —", "Иванов И.И.", "Петров П.П.", "Сидорова А.А."]],
                                value="— Все —", width=140)
    report_cat = ft.Dropdown(label="Категория", options=[ft.dropdown.Option("— Все —")] + [ft.dropdown.Option(x) for x in
                                                                                       ["Свадебные", "Женские",
                                                                                        "Мужские", "Детские"]],
                             value="— Все —", width=140)
    report_mat = ft.Dropdown(label="Материал", options=[ft.dropdown.Option("— Все —")] + [ft.dropdown.Option(x) for x in
                                                                                      ["Золото", "Серебро"]],
                             value="— Все —", width=140)

    report_table = ft.DataTable(
        columns=[
            ft.DataColumn(ft.Text("ID")), ft.DataColumn(ft.Text("Дата")),
            ft.DataColumn(ft.Text("Продавец")), ft.DataColumn(ft.Text("Товар")),
            ft.DataColumn(ft.Text("Мат.")), ft.DataColumn(ft.Text("Кат.")),
            ft.DataColumn(ft.Text("Кол-во")), ft.DataColumn(ft.Text("Цена")),
            ft.DataColumn(ft.Text("Оплата")), ft.DataColumn(ft.Text("Итого"))
        ],
        rows=[], column_spacing=15, heading_row_height=40, data_row_min_height=35
    )
    report_summary = ft.Text("", size=14, weight="bold")

    def load_reports(popular=False):
        report_table.rows.clear()
        start = report_date_from.value
        end = report_date_to.value
        seller = report_seller.value
        category = report_cat.value
        material = report_mat.value
        conn = connect_to_db()
        if not conn: return
        try:
            cursor = conn.cursor(pymysql.cursors.DictCursor)
            if popular:
                query = """
                    SELECT CONCAT(MAX(j.prefix), '-', MIN(j.id)) as full_id,
                           DATE_FORMAT(MAX(s.sale_date), '%Y-%m-%d') as sale_date,
                           '—' as seller_name, s.item_name, j.metal, j.category,
                           SUM(s.quantity_sold) as total_qty, AVG(s.price_per_item) as price,
                           '—' as payment_method, SUM(s.total_price) as total_price
                    FROM sales_history s
                    JOIN jewelry_items j ON s.item_name = j.name
                    WHERE 1=1
                """
                params = []
                if start: query += " AND s.sale_date >= %s"; params.append(start)
                if end: query += " AND s.sale_date <= %s"; params.append(end)
                query += " GROUP BY j.id, s.item_name, j.metal, j.category ORDER BY total_qty DESC"
            else:
                query = """
                    SELECT s.id, s.sale_date, s.seller_name, s.item_name, j.metal, j.category,
                           s.quantity_sold, s.price_per_item, s.payment_method, s.total_price
                    FROM sales_history s
                    LEFT JOIN jewelry_items j ON s.item_name = j.name
                    WHERE 1=1
                """
                params = []
                if start: query += " AND s.sale_date >= %s"; params.append(start)
                if end: query += " AND s.sale_date <= %s"; params.append(end)
                if seller != "— Все —": query += " AND s.seller_name = %s"; params.append(seller)
                if category != "— Все —": query += " AND j.category = %s"; params.append(category)
                if material != "— Все —": query += " AND j.metal = %s"; params.append(material)
                query += " ORDER BY s.sale_date DESC"

            cursor.execute(query, params)
            rows = cursor.fetchall()
            total_sum = total_qty = 0
            for r in rows:
                metal_val = r.get('metal')
                cat_val = r.get('category')
                metal_display = metal_val if metal_val else "—"
                cat_display = cat_val if cat_val else "—"

                report_table.rows.append(ft.DataRow(cells=[
                    ft.DataCell(ft.Text(str(r.get('full_id', r.get('id'))))),
                    ft.DataCell(ft.Text(str(r.get('sale_date', '')))),
                    ft.DataCell(ft.Text(str(r.get('seller_name', '')))),
                    ft.DataCell(ft.Text(str(r.get('item_name', '')))),
                    ft.DataCell(ft.Text(metal_display)),
                    ft.DataCell(ft.Text(cat_display)),
                    ft.DataCell(ft.Text(str(r.get('quantity_sold', r.get('total_qty', 0))))),
                    ft.DataCell(ft.Text(f"{r.get('price_per_item', r.get('price', 0)):.2f}")),
                    ft.DataCell(ft.Text(str(r.get('payment_method', '—')))),
                    ft.DataCell(ft.Text(f"{r.get('total_price', 0):.2f}"))
                ]))
                total_sum += r.get('total_price', 0)
                total_qty += r.get('quantity_sold', r.get('total_qty', 0))
            report_summary.value = f"💰 Выручка: {total_sum:,.2f} руб. | 📦 Всего: {total_qty} шт."
            cursor.close()
            conn.close()
        except Exception as ex:
            print(ex)
        page.update()

    def load_efficiency():
        report_table.rows.clear()
        start = report_date_from.value
        end = report_date_to.value
        conn = connect_to_db()
        if not conn: return
        try:
            cursor = conn.cursor(pymysql.cursors.DictCursor)
            query = """
                SELECT seller_name, COUNT(*) as sales_count, SUM(quantity_sold) as total_qty, SUM(total_price) as total_revenue
                FROM sales_history WHERE 1=1
            """
            params = []
            if start: query += " AND sale_date >= %s"; params.append(start)
            if end: query += " AND sale_date <= %s"; params.append(end)
            query += " GROUP BY seller_name ORDER BY total_revenue DESC"
            cursor.execute(query, params)
            for r in cursor.fetchall():
                report_table.rows.append(ft.DataRow(cells=[
                    ft.DataCell(ft.Text("—")), ft.DataCell(ft.Text("—")),
                    ft.DataCell(ft.Text(r['seller_name'])), ft.DataCell(ft.Text("—")),
                    ft.DataCell(ft.Text("—")), ft.DataCell(ft.Text("—")),
                    ft.DataCell(ft.Text(str(r['total_qty']))), ft.DataCell(ft.Text("—")),
                    ft.DataCell(ft.Text("—")), ft.DataCell(ft.Text(f"{r['total_revenue']:.2f}"))
                ]))
            report_summary.value = "📊 Отчет по эффективности сотрудников"
            cursor.close()
            conn.close()
        except Exception as ex:
            print(ex)
        page.update()

    def save_report_to_docx(e):
        rows_data = [row.cells for row in report_table.rows]
        if not rows_data:
            return
        doc = Document()
        section = doc.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height
        doc.add_heading("Отчет по продажам", 0)
        table = doc.add_table(rows=1, cols=10)
        table.style = "Table Grid"
        widths = [0.5, 0.9, 0.8, 0.9, 0.8, 0.8, 0.7, 0.7, 0.8, 0.8]
        headers = ['ID', 'Дата', 'Продавец', 'Товар', 'Материал', 'Категория', 'Кол-во проданных', 'Цена', 'Оплата',
                   'Итого']
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = h
            table.columns[i].width = Inches(widths[i])
        for row_cells in rows_data:
            r = table.add_row().cells
            for i, cell in enumerate(row_cells):
                r[i].text = str(cell.content.value)
        file_path = os.path.join(tempfile.gettempdir(), "report.docx")
        doc.save(file_path)
        show_message("Отчет успешно сохранен в Word!", "green", target="form")

    def reset_all_report_filters():
        report_date_from.value = ""
        report_date_to.value = ""
        report_seller.value = "— Все —"
        report_cat.value = "— Все —"
        report_mat.value = "— Все —"
        load_reports()

    def show_reports_fullscreen(e=None):
        load_reports(False)
        page.controls.clear()
        page.add(
            ft.Column([
                ft.Row([
                    ft.TextButton("← Назад", on_click=lambda _: switch_to_catalog(None)),
                    ft.Text("Отчёты", size=16, weight="bold")
                ], alignment=ft.MainAxisAlignment.SPACE_BETWEEN),
                ft.Row([
                    ft.Button("Обновить", on_click=lambda _: load_reports(False), bgcolor=btn_color, color="white"),
                    ft.Button("По популярности", on_click=lambda _: load_reports(True), bgcolor=btn_color, color="white"),
                    ft.Button("Сбросить всё", on_click=lambda _: reset_all_report_filters(), bgcolor=danger_color, color="white"),
                ], alignment=ft.MainAxisAlignment.SPACE_EVENLY, wrap=True),
                ft.Row([
                    ft.Button("Эффективность", on_click=lambda _: load_efficiency(), bgcolor=btn_color, color="white"),
                    ft.Button("Сохранить в Word", on_click=save_report_to_docx, bgcolor="#1E90FF", color="white"),
                ], alignment=ft.MainAxisAlignment.SPACE_EVENLY, wrap=True),
                ft.Text("💡 Листайте отчёт влево/вправо", size=11, italic=True, color="grey600"),
                ft.Container(
                    content=ft.Row([report_table], scroll=ft.ScrollMode.ALWAYS),
                    expand=True,
                    border=ft.Border(top=ft.BorderSide(1, "grey300"), bottom=ft.BorderSide(1, "grey300"),
                                     left=ft.BorderSide(1, "grey300"), right=ft.BorderSide(1, "grey300")),
                    border_radius=5,
                    padding=5
                ),
                report_summary
            ], expand=True)
        )
        page.update()

    # ---------- Встроенная информационная панель ----------
    info_panel = ft.Container(visible=False, padding=10, bgcolor="#E8F0E8", border_radius=10)

    def show_info_panel(title_text, content_text):
        info_panel.content = ft.Column([
            ft.Text(title_text, size=18, weight="bold"),
            ft.Text(content_text),
            ft.Button("Закрыть", on_click=lambda _: hide_info_panel())
        ])
        info_panel.visible = True
        confirm_delete_panel.visible = False
        buy_form.visible = False
        page.update()

    def hide_info_panel(e=None):
        info_panel.visible = False
        page.update()

    def show_manual(e):
        show_info_panel(
            "Руководство пользователя",
            "1. Добавление: Заполните поля и нажмите 'Добавить изделие'.\n"
            "2. Поиск: Вводите название в строку поиска для мгновенной фильтрации.\n"
            "3. Покупка: Выберите товар в таблице и нажмите 'Купить изделие'.\n"
            "4. Отчеты: Нажмите кнопку 'Отчеты' для просмотра статистики продаж."
        )

    def show_about(e):
        show_info_panel(
            "О программе",
            "Система учета продаж ювелирных изделий 'Алмазный путь'\n\n"
            "Программный комплекс предназначен для автоматизации розничной торговли: "
            "учета товарных запасов, фиксации продаж и ведения аналитики покупательского спроса.\n\n"
            "Версия: 3.0\n"
            "Технологии: Python (Flet), MySQL (СУБД)."
        )

    def show_developer(e):
        show_info_panel(
            "О разработчике",
            "Исаев Рамин Захид оглы ИС-943"
        )

    info_menu = ft.PopupMenuButton(
        content=ft.Text("Справочная информация"),
        items=[
            ft.PopupMenuItem(content=ft.Text("Руководство пользователя"), on_click=show_manual),
            ft.PopupMenuItem(content=ft.Text("О программе"), on_click=show_about),
            ft.PopupMenuItem(content=ft.Text("О разработчике"), on_click=show_developer),
        ]
    )

    # ---------- Переключение видимости ----------
    table_container = ft.Container(visible=False)

    def show_table(e=None):
        load_data()
        table_container.content = ft.Column([
            ft.Row([
                ft.Button("Назад", on_click=hide_table, bgcolor=btn_color, color="white"),
                ft.Text("Таблица изделий", size=16, weight="bold")
            ]),
            selected_label,
            buy_status_msg,
            ft.Row([
                ft.Button("Купить изделие", on_click=open_buy_form, bgcolor=btn_color, color="white"),
                delete_btn,
            ], alignment=ft.MainAxisAlignment.SPACE_EVENLY),
            ft.Row([
                ft.Button("Нажмите для обновления данных", on_click=lambda _: load_data(), bgcolor=btn_color, color="white"),
            ]),
            confirm_delete_panel,
            buy_form,
            ft.Container(
                content=ft.Row([data_table], scroll=ft.ScrollMode.ALWAYS),
                expand=True,
                border=ft.Border(top=ft.BorderSide(1, "grey300"), bottom=ft.BorderSide(1, "grey300"),
                                 left=ft.BorderSide(1, "grey300"), right=ft.BorderSide(1, "grey300")),
                border_radius=5,
                padding=5
            )
        ])
        table_container.visible = True
        form_container.visible = False
        page.update()

    def hide_table(e=None):
        table_container.visible = False
        form_container.visible = True
        confirm_delete_panel.visible = False
        buy_form.visible = False
        page.update()

    form_container = ft.Container(
        content=ft.Column([
            ft.Row([search_field, sort_button]),
            ft.Row([filter_type, filter_cat, filter_metal, filter_gem, filter_size, filter_purity], wrap=True, spacing=5),
            ft.Button("Сбросить фильтры", on_click=lambda _: (load_filter_values(), load_data()), bgcolor=btn_color, color="white"),
            ft.Divider(),
            ft.Text("Добавление изделия", size=16, weight="bold", color="#4A7C59"),
            ft.Column([
                f_name, f_type, ft.Button("Применить тип изделия", on_click=update_size_options, bgcolor=btn_color, color="white"),
                f_cat, f_metal, f_gem, f_purity, f_weight, f_price, f_size, f_quantity
            ], spacing=7),
            status_msg,
            ft.Row([
                ft.Button("Добавить изделие", on_click=add_item, bgcolor=btn_color, color="white"),
                ft.Button("Открыть таблицу", on_click=show_table, bgcolor=btn_color, color="white"),
            ], alignment=ft.MainAxisAlignment.SPACE_EVENLY),
        ]),
        visible=True
    )

    # ---------- Навигация ----------
    def switch_to_catalog(e):
        page.controls.clear()
        page.add(nav_bar, info_panel, ft.Column([form_container, table_container], expand=True, scroll=ft.ScrollMode.AUTO))
        page.update()

    def switch_to_reports(e):
        page.controls.clear()
        reports_view = ft.Column([
            ft.Text("Фильтры отчётов", size=18, weight="bold"),
            ft.Row([
                report_date_from,
                ft.Button("📅", on_click=lambda _: open_date_picker(report_date_from)),
                report_date_to,
                ft.Button("📅", on_click=lambda _: open_date_picker(report_date_to)),
            ], spacing=5),
            ft.Column([report_seller, report_cat, report_mat], spacing=7),
            ft.Button("Показать собранный отчёт", on_click=show_reports_fullscreen, bgcolor=btn_color, color="white"),
        ], scroll=ft.ScrollMode.AUTO, expand=True)
        page.add(nav_bar, info_panel, reports_view)
        page.update()

    nav_bar = ft.Row([
        ft.Button("Каталог", on_click=switch_to_catalog, bgcolor=btn_color, color="white"),
        ft.Button("Отчеты", on_click=switch_to_reports, bgcolor=btn_color, color="white"),
        info_menu
    ], alignment=ft.MainAxisAlignment.CENTER, spacing=10)

    page.add(nav_bar, info_panel, ft.Column([form_container, table_container], expand=True, scroll=ft.ScrollMode.AUTO))
    load_data()
    page.update()

if __name__ == "__main__":
    ft.run(main)